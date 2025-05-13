from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import json
from pathlib import Path
import sys

# Add project root to Python path
project_root = Path(__file__).parent.parent.parent.parent.parent
sys.path.append(str(project_root))

def generate_management_memo():
    """Generate a management memo report from message data"""
    
    # Initialize document
    doc = Document()
    
    # Add title
    title = doc.add_heading('GPT-4o Usage Analytics Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_paragraph.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    
    # Add executive summary
    doc.add_heading('Executive Summary', level=1)
    summary = doc.add_paragraph()
    summary.add_run('This report provides an analysis of GPT-4o usage patterns and key insights from recent conversations.')
    
    # Load and analyze message data
    messages_dir = project_root / 'storage' / 'datastore' / 'gpt-4o' / 'messages'
    message_files = list(messages_dir.glob('*.json'))
    
    # Add usage statistics
    doc.add_heading('Usage Statistics', level=1)
    stats = doc.add_paragraph()
    stats.add_run(f'Total Conversations: {len(message_files)}')
    
    # Add recent conversations section
    doc.add_heading('Recent Conversations', level=1)
    recent_table = doc.add_table(rows=1, cols=3)
    recent_table.style = 'Table Grid'
    
    # Add table headers
    header_cells = recent_table.rows[0].cells
    header_cells[0].text = 'Date'
    header_cells[1].text = 'Topic'
    header_cells[2].text = 'Key Points'
    
    # Add recent conversations
    for msg_file in sorted(message_files, key=lambda x: x.stat().st_mtime, reverse=True)[:5]:
        with open(msg_file, 'r') as f:
            msg_data = json.load(f)
            row_cells = recent_table.add_row().cells
            row_cells[0].text = datetime.fromtimestamp(msg_file.stat().st_mtime).strftime('%Y-%m-%d')
            row_cells[1].text = msg_data.get('topic', 'N/A')
            row_cells[2].text = msg_data.get('summary', 'N/A')
    
    # Add recommendations section
    doc.add_heading('Recommendations', level=1)
    recommendations = doc.add_paragraph()
    recommendations.add_run('Based on the analysis of conversation data, we recommend:')
    
    # Add bullet points
    doc.add_paragraph('• Continue monitoring usage patterns for optimization opportunities', style='List Bullet')
    doc.add_paragraph('• Consider expanding GPT-4o integration to additional use cases', style='List Bullet')
    doc.add_paragraph('• Implement regular review of conversation quality and outcomes', style='List Bullet')
    
    # Save the document
    output_dir = project_root / 'storage' / 'datastore' / 'analytics' / 'reports'
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / f'gpt4o_analytics_report_{datetime.now().strftime("%Y%m%d")}.docx'
    doc.save(output_file)
    
    return output_file

if __name__ == "__main__":
    report_file = generate_management_memo()
    print(f"Generated management memo: {report_file}")
