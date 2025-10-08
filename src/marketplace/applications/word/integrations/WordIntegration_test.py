import pytest
from src.marketplace.applications.word.integrations.WordIntegration import (
    WordIntegration, 
    WordIntegrationConfiguration
)
from docx.document import Document


@pytest.fixture
def word_config():
    """Fixture that returns a WordIntegrationConfiguration instance"""
    return WordIntegrationConfiguration(
        default_font="Arial",
        default_font_size=12
    )


@pytest.fixture
def word_integration(word_config):
    """Fixture that returns a WordIntegration instance"""
    return WordIntegration(word_config)


def test_word_integration_creation(word_integration):
    """Test WordIntegration instance creation"""
    assert word_integration is not None
    assert word_integration._WordIntegration__configuration.default_font == "Arial"
    assert word_integration._WordIntegration__configuration.default_font_size == 12


def test_create_document(word_integration):
    """Test document creation"""
    doc = word_integration.create_document()
    assert doc is not None
    assert isinstance(doc, Document)


def test_generate_from_html(word_integration):
    """Test HTML to Word conversion"""
    html_content = "<h1>Test Document</h1><p>This is a <strong>test</strong> paragraph.</p>"
    
    doc = word_integration.create_document()
    doc_updated = word_integration.generate_from_html(doc, html_content)
    
    assert doc_updated is not None
    assert isinstance(doc_updated, Document)
    
    # Test document has content
    doc_bytes = word_integration.get_document_bytes(doc_updated)
    assert len(doc_bytes) > 0


def test_generate_from_markdown(word_integration):
    """Test Markdown to Word conversion"""
    markdown_content = """# Test Document

This is a **test** paragraph with *italic* text.

## Subheading

- Bullet point 1
- Bullet point 2

1. Numbered item 1
2. Numbered item 2
"""
    
    doc = word_integration.create_document()
    doc_updated = word_integration.generate_from_markdown(doc, markdown_content)
    
    assert doc_updated is not None
    assert isinstance(doc_updated, Document)
    
    # Test document has content
    doc_bytes = word_integration.get_document_bytes(doc_updated)
    assert len(doc_bytes) > 0


def test_generate_from_text(word_integration):
    """Test plain text to Word conversion"""
    text_content = """Test Document

This is a test paragraph.

• Bullet point 1
• Bullet point 2

1. Numbered item 1
2. Numbered item 2
"""
    
    doc = word_integration.create_document()
    doc_updated = word_integration.generate_from_text(doc, text_content)
    
    assert doc_updated is not None
    assert isinstance(doc_updated, Document)
    
    # Test document has content
    doc_bytes = word_integration.get_document_bytes(doc_updated)
    assert len(doc_bytes) > 0


def test_add_paragraph(word_integration):
    """Test adding paragraphs to document"""
    doc = word_integration.create_document()
    
    # Add regular paragraph
    word_integration.add_paragraph(doc, "This is a test paragraph.")
    
    # Add styled paragraph
    word_integration.add_paragraph(
        doc, 
        "This is a bold paragraph.", 
        bold=True, 
        font_size=14
    )
    
    # Test document has content
    doc_bytes = word_integration.get_document_bytes(doc)
    assert len(doc_bytes) > 0


def test_add_heading(word_integration):
    """Test adding headings to document"""
    doc = word_integration.create_document()
    
    # Add different heading levels
    word_integration.add_heading(doc, "Main Heading", level=1)
    word_integration.add_heading(doc, "Subheading", level=2)
    word_integration.add_heading(doc, "Sub-subheading", level=3)
    
    # Test document has content
    doc_bytes = word_integration.get_document_bytes(doc)
    assert len(doc_bytes) > 0


def test_get_document_bytes(word_integration):
    """Test converting document to bytes"""
    doc = word_integration.create_document()
    word_integration.add_paragraph(doc, "Test content")
    
    doc_bytes = word_integration.get_document_bytes(doc)
    
    assert isinstance(doc_bytes, bytes)
    assert len(doc_bytes) > 0


def test_configuration_defaults():
    """Test default configuration values"""
    config = WordIntegrationConfiguration()
    
    assert config.template_path is None
    assert config.default_font == "Calibri"
    assert config.default_font_size == 11


def test_configuration_custom():
    """Test custom configuration values"""
    config = WordIntegrationConfiguration(
        template_path="/path/to/template.docx",
        default_font="Times New Roman",
        default_font_size=14
    )
    
    assert config.template_path == "/path/to/template.docx"
    assert config.default_font == "Times New Roman"
    assert config.default_font_size == 14
