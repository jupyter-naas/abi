from lib.abi.integration.integration import Integration, IntegrationConfiguration
from dataclasses import dataclass
from typing import List, Optional, Dict, Any
from docx import Document
from docx.document import Document as DocumentType
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import io
import re
import markdown
from pathlib import Path


@dataclass
class WordIntegrationConfiguration(IntegrationConfiguration):
    """Configuration for Word integration.

    Attributes:
        template_path (str, optional): Path to Word template file (.docx)
        default_font (str): Default font family for the document
        default_font_size (int): Default font size in points
    """

    template_path: Optional[str] = None
    default_font: str = "Calibri"
    default_font_size: int = 11


class WordIntegration(Integration):
    """Word integration for creating and modifying documents from templates.

    This class provides methods to interact with Word documents using python-docx,
    allowing creation and modification of documents from markdown/text templates.

    Attributes:
        __configuration (WordIntegrationConfiguration): Configuration instance
            containing template and formatting settings.

    Example:
        >>> config = WordIntegrationConfiguration(
        ...     template_path="path/to/template.docx",
        ...     default_font="Arial",
        ...     default_font_size=12
        ... )
        >>> integration = WordIntegration(config)
    """

    __configuration: WordIntegrationConfiguration
    __document: Optional[DocumentType] = None

    def __init__(self, configuration: WordIntegrationConfiguration):
        super().__init__(configuration)
        self.__configuration = configuration

    def create_document(self, template_path: Optional[str] = None) -> DocumentType:
        """Create a new document from template or blank.

        Args:
            template_path (str, optional): Path to Word template file

        Returns:
            Document: New Word document object
        """
        template_path = template_path or self.__configuration.template_path
        if template_path and Path(template_path).exists():
            self.__document = Document(template_path)
        else:
            self.__document = Document()
            self._setup_default_styles()
        return self.__document

    def _setup_default_styles(self) -> None:
        """Setup default styles for the document."""
        if not self.__document:
            return

        # Configure Normal style
        normal_style = self.__document.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = self.__configuration.default_font
        normal_font.size = Pt(self.__configuration.default_font_size)

        # Create custom styles if they don't exist
        styles = self.__document.styles
        
        # Title style
        if 'Custom Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = self.__configuration.default_font
            title_font.size = Pt(18)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)

        # Heading styles
        for level in range(1, 4):
            style_name = f'Custom Heading {level}'
            if style_name not in [s.name for s in styles]:
                heading_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                heading_font = heading_style.font
                heading_font.name = self.__configuration.default_font
                heading_font.size = Pt(16 - level * 2)
                heading_font.bold = True
                heading_style.paragraph_format.space_before = Pt(12)
                heading_style.paragraph_format.space_after = Pt(6)

    def save_document(self, document: DocumentType, output_path: str) -> None:
        """Save the document to a file.

        Args:
            document (Document): Word document object
            output_path (str): Path to save document to

        Example:
            >>> integration.save_document(doc, "output.docx")
        """
        document.save(output_path)

    def generate_from_markdown(self, document: DocumentType, markdown_content: str) -> DocumentType:
        """Generate document content from markdown input.

        Args:
            document (Document): Word document object
            markdown_content (str): Markdown content to convert

        Returns:
            Document: Updated document with markdown content
        """
        # Parse markdown to HTML first
        html_content = markdown.markdown(markdown_content, extensions=['tables', 'fenced_code'])
        
        # Convert HTML to Word content
        return self._html_to_word(document, html_content)

    def generate_from_html(self, document: DocumentType, html_content: str) -> DocumentType:
        """Generate document content from HTML input.

        Args:
            document (Document): Word document object
            html_content (str): HTML content to convert

        Returns:
            Document: Updated document with HTML content converted to Word format
        """
        return self._html_to_word(document, html_content)

    def generate_from_text(self, document: DocumentType, text_content: str, 
                          preserve_formatting: bool = True) -> DocumentType:
        """Generate document content from plain text input.

        Args:
            document (Document): Word document object
            text_content (str): Plain text content
            preserve_formatting (bool): Whether to preserve line breaks and spacing

        Returns:
            Document: Updated document with text content
        """
        lines = text_content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Add empty paragraph for blank lines
                if preserve_formatting:
                    document.add_paragraph()
                continue
            
            # Detect headings (lines starting with #)
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                heading_text = line.lstrip('# ').strip()
                if level <= 3:
                    style_name = f'Custom Heading {level}'
                    if style_name in [s.name for s in document.styles]:
                        document.add_paragraph(heading_text, style=style_name)
                    else:
                        p = document.add_paragraph(heading_text)
                        p.style = document.styles[f'Heading {min(level, 3)}']
                else:
                    document.add_paragraph(heading_text)
            
            # Detect bullet points
            elif line.startswith(('- ', '* ', '+ ')):
                bullet_text = line[2:].strip()
                p = document.add_paragraph(bullet_text, style='List Bullet')
            
            # Detect numbered lists
            elif re.match(r'^\d+\.\s', line):
                numbered_text = re.sub(r'^\d+\.\s', '', line)
                p = document.add_paragraph(numbered_text, style='List Number')
            
            # Regular paragraph
            else:
                document.add_paragraph(line)
        
        return document

    def _html_to_word(self, document: DocumentType, html_content: str) -> DocumentType:
        """Convert HTML content to Word document elements.

        Args:
            document (Document): Word document object
            html_content (str): HTML content to convert

        Returns:
            Document: Updated document
        """
        # Simple HTML to Word conversion
        # Remove HTML tags and convert to plain text for now
        # This is a simplified implementation - could be enhanced with proper HTML parsing
        
        from html.parser import HTMLParser
        
        class WordHTMLParser(HTMLParser):
            def __init__(self, doc):
                super().__init__()
                self.doc = doc
                self.current_text = ""
                self.in_heading = False
                self.heading_level = 1
                self.in_paragraph = False
                self.in_list = False
                self.in_table = False
                self.table_data = []
                self.current_row = []
                self.in_cell = False
                self.skip_tags = {'style', 'script', 'head', 'meta', 'title'}
                self.current_tag = None
                
            def handle_starttag(self, tag, attrs):
                self.current_tag = tag
                
                if tag in self.skip_tags:
                    return
                    
                if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    self.in_heading = True
                    self.heading_level = int(tag[1])
                elif tag == 'p':
                    self.in_paragraph = True
                elif tag in ['ul', 'ol']:
                    self.in_list = True
                elif tag == 'li':
                    pass  # Handle in handle_data
                elif tag == 'br':
                    self.current_text += '\n'
                elif tag == 'table':
                    self.in_table = True
                    self.table_data = []
                elif tag == 'tr' and self.in_table:
                    self.current_row = []
                elif tag in ['td', 'th'] and self.in_table:
                    self.in_cell = True
                elif tag in ['strong', 'b']:
                    pass  # Handle bold formatting in text
                elif tag in ['em', 'i']:
                    pass  # Handle italic formatting in text
                elif tag == 'div':
                    # Check for special div classes
                    class_attr = dict(attrs).get('class', '')
                    if 'header' in class_attr:
                        pass  # Special header handling
                    
            def handle_endtag(self, tag):
                if tag in self.skip_tags:
                    return
                    
                if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    if self.current_text.strip():
                        style_name = f'Custom Heading {min(self.heading_level, 3)}'
                        if style_name in [s.name for s in self.doc.styles]:
                            self.doc.add_paragraph(self.current_text.strip(), style=style_name)
                        else:
                            p = self.doc.add_paragraph(self.current_text.strip())
                            try:
                                p.style = self.doc.styles[f'Heading {min(self.heading_level, 3)}']
                            except KeyError:
                                # Fallback if heading style doesn't exist
                                run = p.runs[0] if p.runs else p.add_run(self.current_text.strip())
                                run.font.bold = True
                                run.font.size = Pt(16 - self.heading_level * 2)
                    self.current_text = ""
                    self.in_heading = False
                elif tag == 'p':
                    if self.current_text.strip():
                        self.doc.add_paragraph(self.current_text.strip())
                    self.current_text = ""
                    self.in_paragraph = False
                elif tag in ['ul', 'ol']:
                    self.in_list = False
                elif tag == 'li':
                    if self.current_text.strip():
                        self.doc.add_paragraph(self.current_text.strip(), style='List Bullet')
                    self.current_text = ""
                elif tag == 'table':
                    if self.table_data:
                        # Add table to document
                        rows = len(self.table_data)
                        cols = max(len(row) for row in self.table_data) if self.table_data else 0
                        if rows > 0 and cols > 0:
                            table = self.doc.add_table(rows=rows, cols=cols)
                            table.style = 'Table Grid'
                            for row_idx, row_data in enumerate(self.table_data):
                                for col_idx, cell_data in enumerate(row_data):
                                    if row_idx < rows and col_idx < cols:
                                        cell = table.cell(row_idx, col_idx)
                                        cell.text = str(cell_data)
                    self.in_table = False
                    self.table_data = []
                elif tag == 'tr' and self.in_table:
                    if self.current_row:
                        self.table_data.append(self.current_row)
                    self.current_row = []
                elif tag in ['td', 'th'] and self.in_table:
                    if self.current_text.strip():
                        self.current_row.append(self.current_text.strip())
                    else:
                        self.current_row.append("")
                    self.current_text = ""
                    self.in_cell = False
                elif tag == 'div':
                    # Add paragraph break for div elements
                    if self.current_text.strip():
                        self.doc.add_paragraph(self.current_text.strip())
                        self.current_text = ""
                    
            def handle_data(self, data):
                # Skip data in style/script tags
                if self.current_tag in self.skip_tags:
                    return
                    
                # Clean up whitespace but preserve structure
                cleaned_data = data.strip()
                if cleaned_data:
                    if self.current_text and not self.current_text.endswith(' '):
                        self.current_text += ' '
                    self.current_text += cleaned_data
                
            def close(self):
                # Handle any remaining text
                if self.current_text.strip():
                    self.doc.add_paragraph(self.current_text.strip())
                super().close()
        
        parser = WordHTMLParser(document)
        parser.feed(html_content)
        parser.close()
        
        return document

    def add_paragraph(self, document: DocumentType, text: str, 
                     style: Optional[str] = None, alignment: Optional[str] = None,
                     font_size: Optional[int] = None, bold: bool = False,
                     italic: bool = False) -> DocumentType:
        """Add a paragraph to the document with formatting.

        Args:
            document (Document): Word document object
            text (str): Paragraph text
            style (str, optional): Paragraph style name
            alignment (str, optional): Text alignment ('left', 'center', 'right', 'justify')
            font_size (int, optional): Font size in points
            bold (bool): Whether text should be bold
            italic (bool): Whether text should be italic

        Returns:
            Document: Updated document
        """
        p = document.add_paragraph(text, style=style)
        
        # Set alignment
        if alignment:
            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            if alignment.lower() in alignment_map:
                p.alignment = alignment_map[alignment.lower()]
        
        # Set font formatting
        if font_size or bold or italic:
            for run in p.runs:
                if font_size:
                    run.font.size = Pt(font_size)
                if bold:
                    run.font.bold = bold
                if italic:
                    run.font.italic = italic
        
        return document

    def add_heading(self, document: DocumentType, text: str, level: int = 1) -> DocumentType:
        """Add a heading to the document.

        Args:
            document (Document): Word document object
            text (str): Heading text
            level (int): Heading level (1-3)

        Returns:
            Document: Updated document
        """
        level = max(1, min(level, 3))  # Clamp between 1 and 3
        style_name = f'Custom Heading {level}'
        
        if style_name in [s.name for s in document.styles]:
            document.add_paragraph(text, style=style_name)
        else:
            document.add_heading(text, level=level)
        
        return document

    def add_table(self, document: DocumentType, data: List[List[str]], 
                 has_header: bool = True) -> DocumentType:
        """Add a table to the document.

        Args:
            document (Document): Word document object
            data (List[List[str]]): Table data as list of rows
            has_header (bool): Whether first row is header

        Returns:
            Document: Updated document
        """
        if not data:
            return document
        
        rows = len(data)
        cols = len(data[0]) if data else 0
        
        table = document.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        for row_idx, row_data in enumerate(data):
            for col_idx, cell_data in enumerate(row_data):
                if row_idx < rows and col_idx < cols:
                    cell = table.cell(row_idx, col_idx)
                    cell.text = str(cell_data)
                    
                    # Format header row
                    if has_header and row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
        
        return document

    def replace_placeholder(self, document: DocumentType, placeholder: str, 
                           replacement: str) -> DocumentType:
        """Replace placeholder text in the document.

        Args:
            document (Document): Word document object
            placeholder (str): Placeholder text to find (e.g., "{{name}}")
            replacement (str): Text to replace placeholder with

        Returns:
            Document: Updated document
        """
        # Replace in paragraphs
        for paragraph in document.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)
        
        # Replace in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, replacement)
        
        return document

    def get_document_structure(self, document: DocumentType) -> Dict[str, Any]:
        """Get the structure of the document.

        Args:
            document (Document): Word document object

        Returns:
            Dict: Document structure information
        """
        structure: Dict[str, Any] = {
            "paragraphs": [],
            "tables": [],
            "sections": len(document.sections)
        }
        
        # Analyze paragraphs
        for i, paragraph in enumerate(document.paragraphs):
            para_info = {
                "index": i,
                "text": paragraph.text,
                "style": paragraph.style.name if paragraph.style else "Normal",
                "alignment": str(paragraph.alignment) if paragraph.alignment else "LEFT"
            }
            structure["paragraphs"].append(para_info)
        
        # Analyze tables
        for i, table in enumerate(document.tables):
            table_info = {
                "index": i,
                "rows": len(table.rows),
                "cols": len(table.columns) if table.rows else 0,
                "style": table.style.name if table.style else "Table Grid"
            }
            structure["tables"].append(table_info)
        
        return structure

    def get_document_bytes(self, document: DocumentType) -> bytes:
        """Get the document as bytes.

        Args:
            document (Document): Word document object

        Returns:
            bytes: Document file as bytes
        """
        doc_stream = io.BytesIO()
        document.save(doc_stream)
        return doc_stream.getvalue()


def as_tools(configuration: WordIntegrationConfiguration):
    """Convert Word integration into LangChain tools."""
    from langchain_core.tools import StructuredTool
    from pydantic import BaseModel, Field

    integration = WordIntegration(configuration)

    class CreateDocumentSchema(BaseModel):
        template_path: Optional[str] = Field(None, description="Path to Word template file")

    class GenerateFromMarkdownSchema(BaseModel):
        markdown_content: str = Field(..., description="Markdown content to convert to Word document")

    class GenerateFromHtmlSchema(BaseModel):
        html_content: str = Field(..., description="HTML content to convert to Word document")

    class GenerateFromTextSchema(BaseModel):
        text_content: str = Field(..., description="Plain text content to convert to Word document")
        preserve_formatting: bool = Field(True, description="Whether to preserve line breaks and spacing")

    class AddParagraphSchema(BaseModel):
        text: str = Field(..., description="Paragraph text")
        style: Optional[str] = Field(None, description="Paragraph style name")
        alignment: Optional[str] = Field(None, description="Text alignment")
        font_size: Optional[int] = Field(None, description="Font size in points")
        bold: bool = Field(False, description="Whether text should be bold")
        italic: bool = Field(False, description="Whether text should be italic")

    class AddHeadingSchema(BaseModel):
        text: str = Field(..., description="Heading text")
        level: int = Field(1, description="Heading level (1-3)")

    class ReplacePlaceholderSchema(BaseModel):
        placeholder: str = Field(..., description="Placeholder text to find")
        replacement: str = Field(..., description="Text to replace placeholder with")

    class SaveDocumentSchema(BaseModel):
        output_path: str = Field(..., description="Path to save the document")

    def create_document_wrapper(**kwargs):
        doc = integration.create_document(**kwargs)
        # Store document reference - using name mangling for private attribute
        setattr(integration, '_WordIntegration__document', doc)
        return "Document created successfully"

    def generate_from_markdown_wrapper(**kwargs):
        current_doc = getattr(integration, '_WordIntegration__document', None)
        if not current_doc:
            integration.create_document()
            current_doc = getattr(integration, '_WordIntegration__document')
        integration.generate_from_markdown(current_doc, **kwargs)
        return "Content generated from markdown successfully"

    def generate_from_html_wrapper(**kwargs):
        current_doc = getattr(integration, '_WordIntegration__document', None)
        if not current_doc:
            integration.create_document()
            current_doc = getattr(integration, '_WordIntegration__document')
        integration.generate_from_html(current_doc, **kwargs)
        return "Content generated from HTML successfully"

    def generate_from_text_wrapper(**kwargs):
        current_doc = getattr(integration, '_WordIntegration__document', None)
        if not current_doc:
            integration.create_document()
            current_doc = getattr(integration, '_WordIntegration__document')
        integration.generate_from_text(current_doc, **kwargs)
        return "Content generated from text successfully"

    def add_paragraph_wrapper(**kwargs):
        current_doc = getattr(integration, '_WordIntegration__document', None)
        if not current_doc:
            integration.create_document()
            current_doc = getattr(integration, '_WordIntegration__document')
        integration.add_paragraph(current_doc, **kwargs)
        return "Paragraph added successfully"

    def add_heading_wrapper(**kwargs):
        current_doc = getattr(integration, '_WordIntegration__document', None)
        if not current_doc:
            integration.create_document()
            current_doc = getattr(integration, '_WordIntegration__document')
        integration.add_heading(current_doc, **kwargs)
        return "Heading added successfully"

    def replace_placeholder_wrapper(**kwargs):
        current_doc = getattr(integration, '_WordIntegration__document', None)
        if not current_doc:
            return "No document available. Create a document first."
        integration.replace_placeholder(current_doc, **kwargs)
        return "Placeholder replaced successfully"

    def save_document_wrapper(output_path: str):
        current_doc = getattr(integration, '_WordIntegration__document', None)
        if not current_doc:
            return "No document available. Create a document first."
        
        # Ensure the output path uses the proper storage directory
        from pathlib import Path
        storage_dir = Path("storage/datastore/marketplace/applications/word")
        storage_dir.mkdir(parents=True, exist_ok=True)
        
        # If output_path is just a filename, use the storage directory
        output_file = Path(output_path)
        if not output_file.is_absolute() and len(output_file.parts) == 1:
            full_path = storage_dir / output_path
        else:
            full_path = output_file
            
        integration.save_document(current_doc, str(full_path))
        return f"Document saved to {full_path}"

    def get_structure_wrapper():
        current_doc = getattr(integration, '_WordIntegration__document', None)
        if not current_doc:
            return "No document available. Create a document first."
        return integration.get_document_structure(current_doc)

    return [
        StructuredTool(
            name="word_create_document",
            description="Create a new Word document from template or blank",
            func=create_document_wrapper,
            args_schema=CreateDocumentSchema,
        ),
        StructuredTool(
            name="word_generate_from_markdown",
            description="Generate Word document content from markdown input",
            func=generate_from_markdown_wrapper,
            args_schema=GenerateFromMarkdownSchema,
        ),
        StructuredTool(
            name="word_generate_from_html",
            description="Generate Word document content from HTML input",
            func=generate_from_html_wrapper,
            args_schema=GenerateFromHtmlSchema,
        ),
        StructuredTool(
            name="word_generate_from_text",
            description="Generate Word document content from plain text input",
            func=generate_from_text_wrapper,
            args_schema=GenerateFromTextSchema,
        ),
        StructuredTool(
            name="word_add_paragraph",
            description="Add a formatted paragraph to the document",
            func=add_paragraph_wrapper,
            args_schema=AddParagraphSchema,
        ),
        StructuredTool(
            name="word_add_heading",
            description="Add a heading to the document",
            func=add_heading_wrapper,
            args_schema=AddHeadingSchema,
        ),
        StructuredTool(
            name="word_replace_placeholder",
            description="Replace placeholder text in the document with actual content",
            func=replace_placeholder_wrapper,
            args_schema=ReplacePlaceholderSchema,
        ),
        StructuredTool(
            name="word_save_document",
            description="Save the current document to a file",
            func=lambda output_path: save_document_wrapper(output_path),
            args_schema=SaveDocumentSchema,
        ),
        StructuredTool(
            name="word_get_structure",
            description="Get the structure and content overview of the current document",
            func=get_structure_wrapper,
            args_schema=type('StructureSchema', (BaseModel,), {}),
        ),
    ]
