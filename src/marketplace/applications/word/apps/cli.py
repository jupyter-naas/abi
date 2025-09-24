#!/usr/bin/env python3
"""
Word Module CLI Application

A command-line interface for processing various input formats (HTML, Markdown, Text) 
into professional Word documents using templates.

Usage:
    python src/marketplace/applications/word/apps/cli.py --help
"""

import argparse
import sys
import os
import re
from pathlib import Path
from typing import Optional
from dataclasses import dataclass

# Simple Word document processing using python-docx directly
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError as e:
    print(f"❌ Missing dependencies: {e}")
    print("Please install required packages:")
    print("pip install python-docx")
    sys.exit(1)


@dataclass
class WordConfig:
    """Simple configuration for Word processing."""
    template_path: Optional[str] = None
    default_font: str = "Calibri"
    default_font_size: int = 11


class MarkdownProcessor:
    """
    Process markdown content and convert to Word document using proper style mappings.
    
    Markdown to Word Style Mappings:
    ================================
    
    HEADINGS:
    # Main Title          → Heading 1
    ## Section            → Heading 2  
    ### Subsection        → Heading 3
    #### Subsection       → Heading 4
    ##### Details         → Heading 5
    ###### Specifics      → Heading 6
    
    LISTS:
    - Bullet point        → List Bullet (fallback: List Paragraph)
    1. Numbered item      → List Number (fallback: List Paragraph)
    
    REFERENCES:
    ¹ Source citation     → Subtle Reference (fallback: Intense Reference)
    
    PARAGRAPHS:
    **Bold paragraph**    → Strong
    *Italic paragraph*    → Emphasis  
    Regular text          → Normal
    
    SPECIAL:
    ---                   → Page Break
    *Footer metadata*     → Center-aligned italic (for Report prepared by, Date, Classification)
    """
    
    def __init__(self, doc):
        self.doc = doc
        
    def _add_heading(self, text: str, level: int):
        """Add a heading with proper style mapping."""
        # Remove any remaining ** bold markers from the heading text
        clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        
        # Map markdown headings to specific Word styles from template
        style_mapping = {
            1: 'Heading 1',      # # Main Title
            2: 'Heading 2',      # ## Section  
            3: 'Heading 3',      # ### Subsection
            4: 'Heading 4',      # #### Subsection
            5: 'Heading 5',      # ##### Details
            6: 'Heading 6'       # ###### Specifics
        }
        
        try:
            if level in style_mapping:
                style_name = style_mapping[level]
            else:
                # For H7+ use Heading styles
                style_name = f'Heading {min(level, 9)}'
                
            self.doc.add_paragraph(clean_text, style=style_name)
        except:
            # Fallback to manual formatting
            p = self.doc.add_paragraph(clean_text)
            if p.runs:
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(max(12, 18 - level * 1.5))
            else:
                run = p.add_run(clean_text)
                run.font.bold = True
                run.font.size = Pt(max(12, 18 - level * 1.5))
    
    def _add_bullet_point(self, text: str):
        """Add a bullet point with proper list style."""
        clean_text = self._process_inline_formatting(text)
        try:
            # Try List Bullet first, fallback to List Paragraph
            self.doc.add_paragraph(clean_text, style='List Bullet')
        except:
            try:
                self.doc.add_paragraph(clean_text, style='List Paragraph')
            except:
                self.doc.add_paragraph(f"• {clean_text}")
    
    def _add_numbered_item(self, text: str, number: str = None):
        """Add a numbered list item with proper list style."""
        clean_text = self._process_inline_formatting(text)
        
        # If we have a number, preserve it in case Word styles don't handle numbering
        if number:
            formatted_text = f"{number} {clean_text}"
        else:
            formatted_text = clean_text
            
        try:
            # Try List Number first, fallback to List Paragraph
            self.doc.add_paragraph(formatted_text, style='List Number')
        except:
            try:
                self.doc.add_paragraph(formatted_text, style='List Paragraph')
            except:
                # Final fallback with manual numbering
                self.doc.add_paragraph(formatted_text)
    
    def _add_reference(self, text: str):
        """Add a source reference with proper citation style and bookmark."""
        clean_text = self._process_inline_formatting(text)
        
        # Extract the reference number for bookmark
        ref_match = re.match(r'^([¹²³⁴⁵⁶⁷⁸⁹⁰¹⁰¹⁰²⁰³⁰⁴⁰⁵⁰⁶⁰⁷⁰⁸⁰⁹¹¹¹²¹³¹⁴¹⁵¹⁶¹⁷¹⁸¹⁹²²¹²²²³²⁴²⁵²⁶²⁷²⁸²⁹³³¹³²³³³⁴³⁵³⁶³⁷³⁸³⁹⁴⁴¹⁴²⁴³⁴⁴⁴⁵⁴⁶⁴⁷⁴⁸⁴⁹⁵⁵¹⁵²⁵³⁵⁴⁵⁵⁵⁶⁵⁷⁵⁸⁵⁹⁶⁶¹⁶²⁶³⁶⁴⁶⁵⁶⁶⁶⁷⁶⁸⁶⁹⁷⁷¹⁷²⁷³⁷⁴⁷⁵⁷⁶⁷⁷⁷⁸⁷⁹⁸⁸¹⁸²⁸³⁸⁴⁸⁵⁸⁶⁸⁷⁸⁸⁸⁹⁹⁹¹⁹²⁹³⁹⁴⁹⁵⁹⁶⁹⁷⁹⁸⁹⁹]+)', text)
        
        # Just use normal paragraph since the template styles don't exist
        p = self.doc.add_paragraph(clean_text)
        try:
            p.paragraph_format.space_after = Pt(3)
        except:
            pass
        
        # Add bookmark for this reference
        if ref_match:
            ref_number = ref_match.group(1)
            bookmark_name = f"source_{ref_number}"
            self._add_bookmark(p, bookmark_name)
    
    def _add_footer_metadata(self, text: str):
        """Add footer metadata with italic center alignment."""
        clean_text = self._process_inline_formatting(text)
        p = self.doc.add_paragraph(clean_text)
        try:
            # Make it italic and center-aligned for footer formatting
            if p.runs:
                p.runs[0].font.italic = True
            else:
                run = p.add_run(clean_text)
                run.font.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass
    
    def _add_paragraph(self, text: str):
        """Add a regular paragraph with appropriate style and reference hyperlinks."""
        # Check if this is a special paragraph type
        if text.startswith('**') and text.endswith('**'):
            # Strong/Bold paragraph - use Strong style
            clean_text = self._process_inline_formatting(text)
            try:
                p = self.doc.add_paragraph(style='Strong')
                self._process_text_with_references(clean_text, p)
            except:
                p = self.doc.add_paragraph()
                self._process_text_with_references(clean_text, p)
                if p.runs:
                    p.runs[0].font.bold = True
        elif text.startswith('*') and text.endswith('*') and not text.startswith('**'):
            # Emphasis/Italic paragraph - use Emphasis style
            clean_text = self._process_inline_formatting(text)
            try:
                p = self.doc.add_paragraph(style='Emphasis')
                self._process_text_with_references(clean_text, p)
            except:
                p = self.doc.add_paragraph()
                self._process_text_with_references(clean_text, p)
                if p.runs:
                    p.runs[0].font.italic = True
        else:
            # Regular paragraph - use Normal style
            clean_text = self._process_inline_formatting(text)
            try:
                p = self.doc.add_paragraph(style='Normal')
                self._process_text_with_references(clean_text, p)
            except:
                p = self.doc.add_paragraph()
                self._process_text_with_references(clean_text, p)
    
    def _add_page_break(self):
        """Add a page break."""
        self.doc.add_page_break()
    
    def _add_hyperlink(self, paragraph, text, bookmark_name):
        """Add a hyperlink to a paragraph."""
        # Create hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_name)
        
        # Create run element
        run = OxmlElement('w:r')
        
        # Create run properties for blue color
        rPr = OxmlElement('w:rPr')
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')  # Blue color
        rPr.append(color)
        run.append(rPr)
        
        # Create text element
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        run.append(text_elem)
        
        hyperlink.append(run)
        paragraph._p.append(hyperlink)
    
    def _add_bookmark(self, paragraph, bookmark_name):
        """Add a bookmark to a paragraph."""
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), '0')
        bookmark_start.set(qn('w:name'), bookmark_name)
        
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), '0')
        
        paragraph._p.insert(0, bookmark_start)
        paragraph._p.append(bookmark_end)
        
    def process_markdown(self, markdown_content: str):
        """Process markdown content and add to Word document."""
        lines = markdown_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # Handle headings (including nested ones like #### inside paragraphs)
            if line.startswith('#'):
                # Count the number of # symbols
                hash_count = 0
                for char in line:
                    if char == '#':
                        hash_count += 1
                    else:
                        break
                
                # Extract heading text after the # symbols and any spaces
                heading_text = line[hash_count:].strip()
                
                if heading_text:
                    self._add_heading(heading_text, hash_count)
                i += 1
                continue
            
            # Handle horizontal rules
            if line.strip() == '---':
                self._add_page_break()
                i += 1
                continue
            
            # Handle bullet points - map to proper Word list styles
            if line.startswith(('- ', '* ', '+ ')):
                bullet_text = line[2:].strip()
                self._add_bullet_point(bullet_text)
                i += 1
                continue
            
            # Handle numbered lists - map to proper Word numbered list styles
            if re.match(r'^\d+\.\s', line):
                # Extract the number and the text separately
                match = re.match(r'^(\d+\.)\s(.*)$', line)
                if match:
                    number = match.group(1)  # e.g., "1."
                    numbered_text = match.group(2)  # e.g., "**Market Position:** While Big Four..."
                    self._add_numbered_item(numbered_text, number)
                else:
                    # Fallback if regex doesn't match properly
                    numbered_text = re.sub(r'^\d+\.\s', '', line)
                    self._add_numbered_item(numbered_text)
                i += 1
                continue
            
            # Handle numbered references (like in Sources section with superscript numbers)
            if re.match(r'^[¹²³⁴⁵⁶⁷⁸⁹⁰]+\s', line):
                self._add_reference(line.strip())
                i += 1
                continue
            
            # Handle italic metadata lines (like report footer info) - match actual italic syntax
            if re.match(r'^\*(?!.*\*\*)(?:Report prepared by|Date|Classification).*\*\s*$', line.strip()):
                self._add_footer_metadata(line.strip())
                i += 1
                continue
            
            # Check if this is a bold header followed by bullet points 
            # Covers patterns like **Growth Areas:** and **PwC (PricewaterhouseCoopers)**
            if (line.startswith('**') and line.endswith('**') and 
                i < len(lines) - 1 and lines[i + 1].startswith(('- ', '* ', '+ '))):
                # This is a list header - skip it as it will be redundant with the bullet points
                i += 1
                continue
            
            # Handle regular paragraphs (including multi-line)
            paragraph_lines = [line]
            i += 1
            
            # Collect continuation lines for the same paragraph
            while i < len(lines):
                next_line = lines[i].rstrip()
                
                # Stop if we hit an empty line, heading, list item, numbered reference, footer metadata, or horizontal rule
                if (not next_line or 
                    next_line.startswith('#') or 
                    next_line.startswith(('- ', '* ', '+ ')) or
                    re.match(r'^\d+\.\s', next_line) or
                    re.match(r'^[¹²³⁴⁵⁶⁷⁸⁹⁰]+\s', next_line) or
                    re.match(r'^\*(?!.*\*\*)(?:Report prepared by|Date|Classification).*\*\s*$', next_line.strip()) or
                    next_line.strip() == '---'):
                    break
                    
                paragraph_lines.append(next_line)
                i += 1
            
            # Join the paragraph lines and process
            paragraph_text = ' '.join(paragraph_lines).strip()
            if paragraph_text:
                self._add_paragraph(paragraph_text)
    
    def _process_inline_formatting(self, text: str) -> str:
        """Process inline markdown formatting like **bold** and *italic*."""
        # Remove markdown syntax for now - in future could apply actual Word formatting
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove **bold**
        text = re.sub(r'\*(.*?)\*', r'\1', text)      # Remove *italic*
        return text
    
    def _process_text_with_references(self, text: str, paragraph):
        """Process text and convert reference numbers to hyperlinks."""
        # Pattern to match superscript numbers like ¹, ², ³, etc.
        reference_pattern = r'([¹²³⁴⁵⁶⁷⁸⁹⁰¹⁰¹⁰²⁰³⁰⁴⁰⁵⁰⁶⁰⁷⁰⁸⁰⁹¹¹¹²¹³¹⁴¹⁵¹⁶¹⁷¹⁸¹⁹²²¹²²²³²⁴²⁵²⁶²⁷²⁸²⁹³³¹³²³³³⁴³⁵³⁶³⁷³⁸³⁹⁴⁴¹⁴²⁴³⁴⁴⁴⁵⁴⁶⁴⁷⁴⁸⁴⁹⁵⁵¹⁵²⁵³⁵⁴⁵⁵⁵⁶⁵⁷⁵⁸⁵⁹⁶⁶¹⁶²⁶³⁶⁴⁶⁵⁶⁶⁶⁷⁶⁸⁶⁹⁷⁷¹⁷²⁷³⁷⁴⁷⁵⁷⁶⁷⁷⁷⁸⁷⁹⁸⁸¹⁸²⁸³⁸⁴⁸⁵⁸⁶⁸⁷⁸⁸⁸⁹⁹⁹¹⁹²⁹³⁹⁴⁹⁵⁹⁶⁹⁷⁹⁸⁹⁹]+)'
        
        # Split text by reference numbers
        parts = re.split(reference_pattern, text)
        
        for i, part in enumerate(parts):
            if i % 2 == 0:
                # Regular text
                if part:
                    run = paragraph.add_run(part)
            else:
                # Reference number - make it a hyperlink
                bookmark_name = f"source_{part}"
                self._add_hyperlink(paragraph, part, bookmark_name)
        
        return paragraph
    
    def _add_formatted_paragraph(self, text: str):
        """Add a paragraph with proper inline formatting."""
        # For now, just clean the text and add as regular paragraph
        # Future enhancement: could parse **bold** and apply actual formatting
        cleaned_text = self._process_inline_formatting(text)
        self.doc.add_paragraph(cleaned_text)


class WordCLI:
    """Simple command-line interface for Word document generation."""
    
    def __init__(self, config: WordConfig = None):
        self.config = config or WordConfig()
        self.document = None
        
    def create_document(self, template_path: Optional[str] = None):
        """Create a new document."""
        template_path = template_path or self.config.template_path
        
        if template_path and Path(template_path).exists():
            try:
                self.document = Document(template_path)
                print(f"✅ Document created from template: {template_path}")
            except ValueError as e:
                if "is not a Word file" in str(e):
                    # Handle .dotx templates by converting to .docx first
                    print(f"🔄 Converting .dotx template to usable format...")
                    print(f"   Template: {template_path}")
                    
                    # Try to convert .dotx to .docx using a workaround
                    if self._convert_dotx_template(template_path):
                        print(f"✅ Template converted and applied successfully")
                    else:
                        print(f"⚠️  Template conversion failed, using blank document with template styling")
                        self.document = Document()
                        self._setup_custom_styles()
                else:
                    raise e
        else:
            self.document = Document()
            self._setup_default_styles()
            print("✅ Blank document created")
            
    def _convert_dotx_template(self, template_path: str) -> bool:
        """Convert .dotx template to usable document."""
        try:
            import zipfile
            import tempfile
            import shutil
            
            # Create a temporary copy and rename to .docx
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_docx_path = temp_file.name
                
            # Copy the .dotx file to .docx
            shutil.copy2(template_path, temp_docx_path)
            
            # Try to open as .docx
            try:
                self.document = Document(temp_docx_path)
                # Clean up temp file
                Path(temp_docx_path).unlink()
                return True
            except:
                # Clean up temp file
                Path(temp_docx_path).unlink()
                # Fallback to styled blank document
                self.document = Document()
                self._setup_forvis_mazars_styles()
                return False
                
        except Exception as e:
            print(f"   Conversion error: {e}")
            self.document = Document()
            self._setup_forvis_mazars_styles()
            return False
            
    def _setup_default_styles(self):
        """Setup default styles for the document."""
        if not self.document:
            return
        
        # Configure Normal style
        normal_style = self.document.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = self.config.default_font
        normal_font.size = Pt(self.config.default_font_size)
        
    def _setup_custom_styles(self):
        """Setup custom document styles."""
        if not self.document:
            return
            
        # Configure Normal style with standard formatting
        normal_style = self.document.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = "Calibri"
        normal_font.size = Pt(11)
        
        # Add custom header styling if needed
        try:
            # Create custom heading styles if not present in template
            styles = self.document.styles
            if 'Custom Title' not in [s.name for s in styles]:
                title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
                title_font = title_style.font
                title_font.name = "Calibri"
                title_font.size = Pt(24)
                title_font.bold = True
                title_font.color.rgb = (0, 102, 204)  # Professional blue
                title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_style.paragraph_format.space_after = Pt(12)
                
            if 'Custom Subtitle' not in [s.name for s in styles]:
                subtitle_style = styles.add_style('Custom Subtitle', WD_STYLE_TYPE.PARAGRAPH)
                subtitle_font = subtitle_style.font
                subtitle_font.name = "Calibri"
                subtitle_font.size = Pt(16)
                subtitle_font.color.rgb = (51, 51, 51)  # Dark gray
                subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                subtitle_style.paragraph_format.space_after = Pt(6)
                
        except Exception as e:
            print(f"   Style setup warning: {e}")
            # Continue with basic styling
        
    def process_markdown_file(self, input_path: str):
        """Process a Markdown file."""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
            
            if not self.document:
                self.create_document()
                
            # Process markdown directly without HTML conversion
            processor = MarkdownProcessor(self.document)
            processor.process_markdown(markdown_content)
            
            print(f"✅ Processed Markdown file: {input_path}")
            return True
            
        except FileNotFoundError:
            print(f"❌ Error: Markdown file not found: {input_path}")
            return False
        except Exception as e:
            print(f"❌ Error processing Markdown file: {e}")
            return False
            
    def replace_placeholders(self, replacements: dict):
        """Replace placeholders in the document."""
        if not self.document:
            print("❌ Error: No document created. Create a document first.")
            return False
            
        for placeholder, replacement in replacements.items():
            # Replace in paragraphs
            for paragraph in self.document.paragraphs:
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, replacement)
            
            # Replace in tables
            for table in self.document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, replacement)
            
            print(f"✅ Replaced '{placeholder}' with '{replacement}'")
        return True
        
    def save_document(self, output_path: str, input_filename: str = None, add_timestamp: bool = True):
        """Save the document with optional timestamp prefix using input filename."""
        if not self.document:
            print("❌ Error: No document to save. Create and process content first.")
            return False
            
        try:
            # Add timestamp prefix if requested
            if add_timestamp:
                from datetime import datetime
                timestamp = datetime.now().strftime("%Y%m%dT%H%M%S_")
                
                if input_filename:
                    # Use input filename (without extension) with timestamp prefix
                    input_path = Path(input_filename)
                    base_name = input_path.stem  # filename without extension
                    timestamped_name = timestamp + base_name + ".docx"
                else:
                    # Use provided output filename with timestamp prefix
                    output_file = Path(output_path)
                    timestamped_name = timestamp + output_file.name
                
                final_output_path = Path(output_path).parent / timestamped_name
            else:
                final_output_path = Path(output_path)
            
            # Ensure output directory exists
            final_output_path.parent.mkdir(parents=True, exist_ok=True)
            
            self.document.save(str(final_output_path))
            print(f"✅ Document saved: {final_output_path}")
            return True
            
        except Exception as e:
            print(f"❌ Error saving document: {e}")
            return False
            
    def get_document_info(self):
        """Display document information."""
        if not self.document:
            print("❌ No document created")
            return
            
        paragraph_count = len(self.document.paragraphs)
        table_count = len(self.document.tables)
        section_count = len(self.document.sections)
        
        print(f"\n📊 Document Information:")
        print(f"   Paragraphs: {paragraph_count}")
        print(f"   Tables: {table_count}")
        print(f"   Sections: {section_count}")
        
        if self.document.paragraphs:
            print(f"\n📝 First few paragraphs:")
            for i, para in enumerate(self.document.paragraphs[:3]):
                text_preview = para.text[:50] + "..." if len(para.text) > 50 else para.text
                print(f"   {i+1}. {text_preview}")


def create_parser():
    """Create the argument parser."""
    parser = argparse.ArgumentParser(
        description="Word Module CLI - Generate professional Word documents from Markdown",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Process Markdown file with template
  python cli.py --markdown content.md --template template.docx --output document.docx
  
  # Process with placeholder replacement
  python cli.py --markdown content.md --replace "{{name}}" "John Doe" --output report.docx
  
  # Process Markdown report with template
  python cli.py --markdown company_report.md --template company-template.docx --output company_report.docx
        """
    )
    
    # Input - only markdown
    parser.add_argument('--markdown', '-m', required=True, help='Markdown file to process')
    
    # Template and output
    parser.add_argument('--template', help='Word template file (.docx)')
    parser.add_argument('--output', '-o', required=True, help='Output Word document path')
    
    # Placeholder replacement
    parser.add_argument('--replace', nargs=2, metavar=('PLACEHOLDER', 'VALUE'), 
                       action='append', help='Replace placeholder with value (can be used multiple times)')
    
    # Options
    parser.add_argument('--info', action='store_true', help='Show document information after processing')
    parser.add_argument('--verbose', '-v', action='store_true', help='Verbose output')
    parser.add_argument('--no-timestamp', action='store_true', help='Disable timestamp prefix on output file')
    
    return parser


def main():
    """Main CLI function."""
    parser = create_parser()
    args = parser.parse_args()
    
    # Initialize CLI
    config = WordConfig(template_path=args.template)
    cli = WordCLI(config)
    
    if args.verbose:
        print("🔤 Word Module CLI - Markdown to Word")
        print("=" * 35)
    
    # Create document
    cli.create_document(args.template)
    
    # Process markdown input
    success = cli.process_markdown_file(args.markdown)
    
    if not success:
        sys.exit(1)
    
    # Replace placeholders
    if args.replace:
        replacements = {placeholder: value for placeholder, value in args.replace}
        cli.replace_placeholders(replacements)
    
    # Show info if requested
    if args.info:
        cli.get_document_info()
    
    # Save document with timestamp control
    add_timestamp = not args.no_timestamp
    if not cli.save_document(args.output, args.markdown, add_timestamp):
        sys.exit(1)
    
    if args.verbose:
        print("✅ Processing completed successfully!")


if __name__ == "__main__":
    main()